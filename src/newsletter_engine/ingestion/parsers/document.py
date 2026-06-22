"""Document parser: .pdf via pypdf (page provenance) and generic .docx via python-docx
(heading provenance)."""

from __future__ import annotations

import docx
from pypdf import PdfReader

from newsletter_engine.ingestion.chunker import ParsedChunk


def parse_pdf(path) -> list[ParsedChunk]:
    reader = PdfReader(str(path))
    chunks: list[ParsedChunk] = []
    for number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.append(ParsedChunk(text=text, location={"page": number}))
    return chunks


def parse_docx(path) -> list[ParsedChunk]:
    """Generic (non-transcript) .docx: paragraphs grouped under their nearest heading."""
    document = docx.Document(str(path))
    chunks: list[ParsedChunk] = []
    heading = "Document"
    start_index = 1
    buffer: list[str] = []

    def flush(end_index: int):
        nonlocal buffer, start_index
        if buffer:
            chunks.append(
                ParsedChunk(
                    text="\n".join(buffer),
                    location={"heading": heading, "paragraphs": f"{start_index}-{end_index}"},
                )
            )
            buffer = []
        start_index = end_index + 1

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style is not None and paragraph.style.name.startswith("Heading"):
            flush(index - 1)
            heading = text
            start_index = index
        else:
            buffer.append(text)

    flush(len(document.paragraphs))
    return chunks
